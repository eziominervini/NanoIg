#!/usr/bin/python

import sys, os
import pandas as pd
from selenium.webdriver import Firefox
from selenium.webdriver.firefox.options import Options
from docx import Document
from docx.shared import Cm
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


run=__file__
path=sys.argv[1] 

def pypath(folder):
    pyfolder=''
    pyfolder=folder.replace('\\','/')
    pyfolder=pyfolder.replace('\\\\','/')
    
    return pyfolder

upper=pypath(run).replace("/NanoIg.py", "")

BClist = [f for f in os.listdir(pypath(path)+"/fastq")]

#Interrogate IMGT/V-Quest

fileObject = open(pypath(path) + "/Results.fasta", "r")
consensus = fileObject.read()
print(consensus)
fileObject.close()

opts = Options()
opts.set_headless()
assert opts.headless 
browser = Firefox(options=opts)
browser.get('http://www.imgt.org/IMGT_vquest/analysis')

browser.find_element_by_xpath("//select[@name='species']/option[text()='Homo sapiens (human)']").click()
browser.find_element_by_xpath("//select[@name='receptorOrLocusType']/option[text()='IG']").click()
browser.find_element_by_id("main_form_sequences").click()
browser.find_element_by_id("main_form_sequences").send_keys(consensus)
browser.find_element_by_id("sv_resultType").click()
browser.find_element_by_id("text_sv").click()
browser.find_element_by_id("main_form_V_REGIONsearchIndeltrue").click()
browser.find_element_by_xpath("//select[@name='sv_V_GENEordertable']/option[text()='input']").click()
browser.find_element_by_id("main_form_0").click()

results = browser.find_element_by_tag_name('pre')
print(results.text)

final = open(pypath(path) + "/V_quest.txt", "w")
final.write(results.text)
final.close()

browser.close()

#Write Report

document = Document()

sections = document.sections
section = sections[0]
new_width, new_height = section.page_height, section.page_width
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width = new_width
section.page_height = new_height
section.top_margin=Cm(0.5)
section.bottom_margin=Cm(0.5)
section.left_margin=Cm(1.27)
section.right_margin=Cm(1.27)

document.add_heading('IGHV Mutational Analysis Summary', level=0)

document.add_heading('Clonality Analysis', level=1)

my_image = document.add_picture(pypath(path) + '/PIPE/Clonality/grid.jpg', width=Cm(17))

last_paragraph = document.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

document.add_heading('IMGT V-quest Analysis', level=1)



with open(pypath(path)+"/V_quest.txt", 'r') as infile, open(pypath(path)+"/Summary_txt.txt", "a") as outfile: 
    for line in infile: 
            if "Sequence Number" in line:
                outfile.write(line)

with open(pypath(path)+"/V_quest.txt", 'r') as infile, open(pypath(path)+"/Summary_txt.txt", "a") as outfile: 
    for line in infile: 
            if ";barcode" in line:
                outfile.write(line)

Summary = pd.read_csv(pypath(path) + "/Summary_txt.txt", sep=";", index_col="Sequence Number")

subsummary= Summary[["Sequence ID", "V-GENE and allele", "V-DOMAIN Functionality", "V-REGION identity % (nt)", "J-GENE and allele", "D-GENE and allele", "CDR-IMGT lengths", "AA JUNCTION", "Sequence analysis category"]]
newcol=[]
oldcol=[]
d={}
for i in subsummary["Sequence ID"]:
    newcol.append(i[:i.rindex(':')] + '')
for i in subsummary["Sequence ID"]:
    oldcol.append(i)
for i in range(len(oldcol)):
    d[oldcol[i]]=newcol[i]
subsummary["Sequence ID"].replace(d, inplace=True)

t = document.add_table(subsummary.shape[0]+1, subsummary.shape[1])

for j in range(subsummary.shape[-1]):
    t.cell(0,j).text = subsummary.columns[j]
    
for i in range(subsummary.shape[0]):
    for j in range(subsummary.shape[-1]):
        t.cell(i+1,j).text = str(subsummary.values[i,j])
        
for row in t.rows:
    for cell in row.cells:
        paragraphs = cell.paragraphs
        for paragraph in paragraphs:
            for run in paragraph.runs:
                font = run.font
                font.name='Arial'
                font.size= Pt(7)


document.save(pypath(path) + '/Summary_report.docx')

quit()
